"""md_to_docx.py - convert findings_report.md into findings_report.docx,
stripping the `---` separator lines between sections and embedding the 5
referenced PNG figures inline.

No external pandoc dependency. Uses python-docx directly.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path('findings_report.md')
OUT = Path('findings_report.docx')

text = SRC.read_text(encoding='utf-8')
lines = text.split('\n')

# Strip lines that are just `---` (horizontal-rule separators between sections).
# A bare `---` on its own line is the separator. Keep lines that contain `---`
# inside table syntax (which never look like just `---` on its own).
lines = [ln for ln in lines if ln.strip() != '---']

# Re-join. We may now have triple-blank-line runs; collapse to at most one blank.
out_lines = []
prev_blank = False
for ln in lines:
    is_blank = (ln.strip() == '')
    if is_blank and prev_blank:
        continue
    out_lines.append(ln)
    prev_blank = is_blank
text = '\n'.join(out_lines)
lines = text.split('\n')

doc = Document()
# Default style tweaks for legibility
for style_name in ('Normal',):
    s = doc.styles[style_name]
    s.font.name = 'Calibri'
    s.font.size = Pt(11)


def add_runs_with_inline(paragraph, text):
    """Add text to paragraph with inline bold (**...**) and italic (*...*)
    and inline-code (`...`) handling. Strips them and applies formatting."""
    # Tokenize via regex. Order matters: code first (because backticks can contain * etc).
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def add_paragraph(text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs_with_inline(p, text)
    return p


def looks_like_table_row(s):
    return s.lstrip().startswith('|') and s.rstrip().endswith('|')


def is_table_separator_row(s):
    # |---|---|---:|---:|  shape: pipes around dashes/colons/spaces only
    inner = s.strip()
    if not (inner.startswith('|') and inner.endswith('|')):
        return False
    inner = inner.strip('|')
    cells = inner.split('|')
    return all(re.fullmatch(r'\s*:?-+:?\s*', c) for c in cells)


def parse_table_row(s):
    inner = s.strip().strip('|')
    return [c.strip() for c in inner.split('|')]


# Walk the lines and emit doc content.
i = 0
N = len(lines)
while i < N:
    ln = lines[i]
    stripped = ln.strip()

    # Headings
    if stripped.startswith('# '):
        h = doc.add_heading('', level=0)
        r = h.add_run(stripped[2:].strip())
        i += 1
        continue
    if stripped.startswith('## '):
        h = doc.add_heading('', level=1)
        r = h.add_run(stripped[3:].strip())
        i += 1
        continue
    if stripped.startswith('### '):
        h = doc.add_heading('', level=2)
        r = h.add_run(stripped[4:].strip())
        i += 1
        continue

    # Image: ![alt](path) optionally followed by an italic caption line
    m_img = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
    if m_img:
        alt, path = m_img.group(1), m_img.group(2)
        img_path = Path(path)
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(6.3))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = doc.add_paragraph()
                p.add_run(f'[image could not be embedded: {path} ({e})]').italic = True
        else:
            p = doc.add_paragraph()
            p.add_run(f'[image not found at: {path}]').italic = True
        # Look ahead for caption line (next non-blank line starting with *)
        i += 1
        # Skip blank lines
        while i < N and lines[i].strip() == '':
            i += 1
        if i < N and lines[i].strip().startswith('*') and lines[i].strip().endswith('*') and not lines[i].strip().startswith('**'):
            cap = lines[i].strip().strip('*').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(cap); r.italic = True; r.font.size = Pt(10)
            i += 1
        continue

    # Tables
    if looks_like_table_row(ln) and i + 1 < N and is_table_separator_row(lines[i+1]):
        header = parse_table_row(ln)
        i += 2  # past header + separator
        rows = []
        while i < N and looks_like_table_row(lines[i]):
            rows.append(parse_table_row(lines[i]))
            i += 1
        # Create docx table
        t = doc.add_table(rows=1 + len(rows), cols=len(header))
        t.style = 'Light Grid Accent 1'
        for c, txt in enumerate(header):
            cell = t.rows[0].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt); r.bold = True; r.font.size = Pt(10)
        for ri, row in enumerate(rows, start=1):
            for c, txt in enumerate(row[:len(header)]):
                cell = t.rows[ri].cells[c]
                cell.text = ''
                p = cell.paragraphs[0]
                add_runs_with_inline(p, txt)
                for run in p.runs:
                    run.font.size = Pt(10)
        # Add a small spacer paragraph after the table
        doc.add_paragraph('')
        continue

    # Empty line -> nothing, but skip
    if stripped == '':
        i += 1
        continue

    # Bold-only paragraph leader like **Student:** ... or **AdaptiveFedAvg.** ...
    # Just treat as a normal paragraph; the inline parser will handle **...**.
    # Bullet list (lines starting with - or *)
    if stripped.startswith('- ') or (stripped.startswith('* ') and not stripped.startswith('**')):
        body = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        add_runs_with_inline(p, body)
        i += 1
        continue

    # Plain paragraph: collect contiguous non-empty non-special lines
    para_lines = [ln]
    j = i + 1
    while j < N:
        nxt = lines[j]
        nstr = nxt.strip()
        if nstr == '': break
        if nstr.startswith('#'): break
        if looks_like_table_row(nxt): break
        if re.match(r'!\[', nstr): break
        if nstr.startswith('- '): break
        para_lines.append(nxt)
        j += 1
    body = ' '.join(para_lines).strip()
    if body:
        add_paragraph(body)
    i = j

doc.save(str(OUT))
print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")
